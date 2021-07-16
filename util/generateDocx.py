from docx import Document
import subprocess, time

def generateReport(analysis):

    date = time.strftime("%d-%m-%Y-%H:%M:%S")

    document = Document()

    document.add_heading('Oncogeriatric synthetic data ', 0)

    document.add_paragraph(f'Date: {date}')

    document.add_paragraph(f'Sample size: {analysis.sampleSize()}')

    document.add_heading('Patient characteristics', level=1)

    records = (
        ('Age (years)', f"{round(analysis.ageStats()[0])} ({analysis.ageStats()[2][0]}-{analysis.ageStats()[2][1]})"),
        ('Gender (% female)', f"{round(analysis.genderBalance()*100)}"),
        ('Cancer (%)', '')
    )

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Summary statistics'
    for param, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = value
    for param,value in analysis.cancerPrevalence().items():
        row_cells = table.add_row().cells
        row_cells[0].text = '        '+param.capitalize()
        row_cells[1].text = str(round(value*100, 1))
    row_cells = table.add_row().cells
    row_cells[0].text = 'Anthropometrics'
    row_cells[1].text = ''
    for param,value in analysis.anthropometrics():
        row_cells = table.add_row().cells
        row_cells[0].text = '        '+param
        row_cells[1].text = str(round(value, 1))
    row_cells = table.add_row().cells
    row_cells[0].text = 'Other'
    row_cells[1].text = ''
    row_cells = table.add_row().cells
    row_cells[0].text = f'        Creatinine (mean, μmol/L)'
    row_cells[1].text = str(round(analysis.creatinineMean(), 1))
    row_cells = table.add_row().cells
    row_cells[0].text = f'        Self-reported health (mean, 1-4)'
    row_cells[1].text = str(round(analysis.srhMean(), 1))
    row_cells = table.add_row().cells
    row_cells[0].text = f'        Timed-up-and-Go test (mean, s)'
    row_cells[1].text = str(round(analysis.tugMean(), 1))
    row_cells = table.add_row().cells
    row_cells[0].text = f'        Frailty (mean, %)'
    row_cells[1].text = str(round(analysis.frailtyMean(), 1))
    row_cells = table.add_row().cells
    row_cells[0].text = 'Electronic frailty index'
    row_cells[1].text = ''
    for param,value in analysis.efiMeans().items():
        row_cells = table.add_row().cells
        row_cells[0].text = '        '+param.replace("_", " ").capitalize()
        row_cells[1].text = str(round(value, 1))

    efi_accuracy = analysis.efiAccuracy()

    document.add_paragraph(
        f"The eFI exhibited high sensitivity ({round(efi_accuracy['sensitivity'], 1)}%) for frailty compared to the gold standard Fried’s criteria, but low specificity ({round(efi_accuracy['specificity'], 1)}%) was observed with high negative predictive value ({round(efi_accuracy['negative_predictive_value'], 1)}%), meaning that an eFI below 0.13 was able to reliably exclude frailty in this synthetic oncogeriatric population."
    )

    tug_accuracy = analysis.tugAnalysis()

    document.add_paragraph(
        f"As expected from the reverse engineering of the TUG test diagnostic accuracy from an existing meta-analysis, this generally mirrored that expected with high sensitivity ({round(tug_accuracy['sensitivity'], 1)}%), reasonable specificity ({round(tug_accuracy['specificity'], 1)}%) and high negative predictive value ({round(tug_accuracy['negative_predictive_value'], 1)}%)."
    )

    document.add_page_break()

    analysis.plotExpectedPrevalence()
    document.add_picture('results/plots/prevalence_expected.png')
    
    analysis.FDM()
    document.add_heading('Overlap between frailty, disability and multimorbidity', level=1)
    document.add_picture('results/plots/fdm.png')

    document.add_heading('Relationship between body mass index categories and activity', level=1)
    analysis.bmiActivityPlot()
    document.add_picture('results/plots/funnel.png')

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Outcome'
    hdr_cells[1].text = '% (95% CI)'
    for k,v in analysis.outcomeStats().items():
        row_cells = table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = v

    document.save(f'results/reports/{date}.docx')

    subprocess.call(('open', f'results/reports/{date}.docx'))
